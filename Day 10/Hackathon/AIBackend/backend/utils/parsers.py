from PyPDF2 import PdfReader
import docx
import re
import json

def parse_offer_file(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    elif file_path.endswith(".txt"):
        with open(file_path, "r") as f:
            return f.read()
    else:
        return ""
    
def parse_resume_file(file_path: str) -> str:
    if file_path.endswith(".pdf"):
        reader = PdfReader(file_path)
        return "\n".join([page.extract_text() or "" for page in reader.pages])
    elif file_path.endswith(".docx"):
        doc = docx.Document(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    else:
        with open(file_path, "r") as f:
            return f.read()

def parse_gemini_response(response):
    """Parse Gemini response into a structured dict."""
    
    # If it's already a dict, return as is
    if isinstance(response, dict):
        return response

    # If it's an AIMessage or string
    if hasattr(response, "text") and callable(response.text):
        text = response.text()
    elif hasattr(response, "text"):
        text = response.text
    elif hasattr(response, "content"):
        text = response.content
    else:
        text = str(response)

    try:
        # Remove backticks and labels like ```json
        cleaned = re.sub(r"```(?:json|python)?", "", text).strip("`\n ")

        # If starts with assignment like `x = {...}`, keep only {...}
        if "=" in cleaned and "{" in cleaned:
            cleaned = cleaned.split("=", 1)[-1]

        return json.loads(cleaned)
    except Exception as e:
        return {"error": str(e), "raw": text}
